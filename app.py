from flask import Flask, render_template, request, jsonify, send_file
import threading
import os
import sys
import json
import re
import time
from pathlib import Path

sys.path.insert(0, os.path.dirname(__file__))

from researcher import research_topic
from writer import write_article
from truth_loop import audit_article
from image_collector import collect_images

app = Flask(__name__)

BASE_DIR = Path(__file__).parent
ARTICLES_DIR = BASE_DIR / "articles"
EXPORTS_DIR = BASE_DIR / "exports"
IMAGES_DIR = BASE_DIR / "images"
ARTICLES_DIR.mkdir(exist_ok=True)
EXPORTS_DIR.mkdir(exist_ok=True)

pipeline_state = {
    "stage": "idle",
    "percent": 0,
    "slug": None,
    "error": None,
    "article_text": None,
    "images": [],
}
_lock = threading.Lock()


def _set_state(**kwargs):
    with _lock:
        for k, v in kwargs.items():
            pipeline_state[k] = v


def _slugify(text):
    text = text.lower().strip()
    text = re.sub(r"[^a-z0-9\s-]", "", text)
    text = re.sub(r"\s+", "-", text)
    return text[:80]


def _article_exists(slug):
    return (ARTICLES_DIR / f"{slug}.md").exists()


def _generate_article(topic, mode, article_type):
    try:
        slug = _slugify(topic)
        timestamp = int(time.time())
        slug = f"{slug}-{timestamp}"
        _set_state(stage="researching", percent=10, slug=slug)

        research = research_topic(topic, mode=mode)
        _set_state(stage="writing", percent=35)

        article = write_article(research, topic, mode=mode)
        _set_state(stage="auditing", percent=55)

        revised_article, audit_report = audit_article(article, topic, mode=mode)
        _set_state(stage="images", percent=75)

        entities = research.get("entities", [])
        images = collect_images(revised_article, entities, slug, topic=topic)
        _set_state(stage="publishing", percent=90)

        md_path = ARTICLES_DIR / f"{slug}.md"
        md_path.write_text(revised_article, encoding="utf-8")

        _set_state(
            stage="done",
            percent=100,
            article_text=revised_article,
            images=images,
            error=None,
        )
    except Exception as e:
        _set_state(stage="error", percent=0, error=str(e))


def _render_html(article_text, slug):
    try:
        import markdown
        from markdown.extensions.fenced_code import FencedCodeExtension
        from markdown.extensions.tables import TableExtension
        html_body = markdown.markdown(
            article_text,
            extensions=[FencedCodeExtension(), TableExtension()],
        )
    except ImportError:
        html_body = f"<pre>{article_text}</pre>"

    title = slug.replace("-", " ").title()
    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{title} - AI Journalist Studio</title>
<style>
body {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif; max-width: 800px; margin: 2em auto; padding: 0 1em; line-height: 1.7; color: #e0e0e0; background: #1a1a2e; }}
img {{ max-width: 100%; height: auto; border-radius: 6px; margin: 1em 0; }}
h1, h2, h3, h4 {{ color: #00d4aa; }}
a {{ color: #00d4aa; }}
blockquote {{ border-left: 4px solid #00d4aa; margin: 1em 0; padding-left: 1em; color: #aaa; }}
code {{ background: #2a2a3e; padding: 2px 6px; border-radius: 3px; }}
pre {{ background: #2a2a3e; padding: 1em; border-radius: 6px; overflow-x: auto; }}
table {{ border-collapse: collapse; width: 100%; margin: 1em 0; }}
th, td {{ border: 1px solid #333; padding: 8px; text-align: left; }}
th {{ background: #2a2a3e; color: #00d4aa; }}
</style>
</head>
<body>
{html_body}
</body>
</html>"""


def _generate_docx(article_text, slug):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    import markdown
    from markdown.extensions.fenced_code import FencedCodeExtension
    from markdown.extensions.tables import TableExtension
    from bs4 import BeautifulSoup

    html_body = markdown.markdown(
        article_text,
        extensions=[FencedCodeExtension(), TableExtension()],
    )
    soup = BeautifulSoup(html_body, "html.parser")

    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0xE0, 0xE0, 0xE0)

    for el in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "ul", "ol", "blockquote", "pre", "table", "hr"]):
        if el.name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(el.name[1])
            heading = doc.add_heading(el.get_text(), level=level)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0x00, 0xD4, 0xAA)
        elif el.name == "p":
            if el.find("img"):
                img = el.find("img")
                src = img.get("src", "")
                alt = img.get("alt", "")
                if src and os.path.exists(src):
                    doc.add_picture(src, width=Inches(5))
                if alt:
                    doc.add_paragraph(alt).italic = True
            else:
                doc.add_paragraph(el.get_text())
        elif el.name == "ul":
            for li in el.find_all("li"):
                doc.add_paragraph(li.get_text(), style="List Bullet")
        elif el.name == "ol":
            for li in el.find_all("li"):
                doc.add_paragraph(li.get_text(), style="List Number")
        elif el.name == "blockquote":
            p = doc.add_paragraph(el.get_text())
            p.italic = True
        elif el.name == "pre":
            code = el.get_text()
            p = doc.add_paragraph(code)
            p.style.font.name = "Consolas"
        elif el.name == "hr":
            doc.add_paragraph("─" * 50)

    path = EXPORTS_DIR / f"{slug}.docx"
    doc.save(str(path))
    return path


def _generate_pdf(article_text, slug):
    html = _render_html(article_text, slug)
    path = EXPORTS_DIR / f"{slug}.pdf"
    path.write_text(html, encoding="utf-8")
    return path


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/generate", methods=["POST"])
def generate():
    data = request.get_json(silent=True) or {}
    topic = data.get("topic", "").strip()
    mode = data.get("mode", "fast")
    article_type = data.get("type", "analytical")

    if not topic:
        return jsonify({"error": "Topic is required"}), 400
    if mode not in ("fast", "deep"):
        return jsonify({"error": "Mode must be 'fast' or 'deep'"}), 400

    if pipeline_state["stage"] not in ("idle", "done", "error"):
        return jsonify({"error": "A generation is already in progress"}), 409

    _set_state(stage="queued", percent=0, slug=None, article_text=None, images=[], error=None)

    thread = threading.Thread(
        target=_generate_article,
        args=(topic, mode, article_type),
        daemon=True,
    )
    thread.start()

    return jsonify({"status": "started"}), 202


@app.route("/progress")
def progress():
    with _lock:
        return jsonify({
            "stage": pipeline_state["stage"],
            "percent": pipeline_state["percent"],
            "slug": pipeline_state["slug"],
            "error": pipeline_state["error"],
        })


@app.route("/article/<slug>")
def article(slug):
    md_path = ARTICLES_DIR / f"{slug}.md"
    if not md_path.exists():
        return jsonify({"error": "Article not found"}), 404
    article_text = md_path.read_text(encoding="utf-8")
    html = _render_html(article_text, slug)
    return html


@app.route("/download/<slug>/<fmt>")
def download(slug, fmt):
    fmt = fmt.lower()
    md_path = ARTICLES_DIR / f"{slug}.md"
    if not md_path.exists():
        return jsonify({"error": "Article not found"}), 404

    article_text = md_path.read_text(encoding="utf-8")

    if fmt == "md":
        return send_file(str(md_path), as_attachment=True, download_name=f"{slug}.md")
    elif fmt == "html":
        html = _render_html(article_text, slug)
        html_path = EXPORTS_DIR / f"{slug}.html"
        html_path.write_text(html, encoding="utf-8")
        return send_file(str(html_path), as_attachment=True, download_name=f"{slug}.html")
    elif fmt == "docx":
        docx_path = _generate_docx(article_text, slug)
        return send_file(str(docx_path), as_attachment=True, download_name=f"{slug}.docx")
    elif fmt == "pdf":
        pdf_path = _generate_pdf(article_text, slug)
        return send_file(
            str(pdf_path),
            as_attachment=True,
            download_name=f"{slug}.pdf",
            mimetype="application/pdf",
        )
    else:
        return jsonify({"error": f"Unsupported format: {fmt}"}), 400


if __name__ == "__main__":
    config_path = BASE_DIR / "config.json"
    if config_path.exists():
        with open(config_path) as f:
            cfg = json.load(f)
        host = cfg.get("web_ui", {}).get("host", "127.0.0.1")
        port = cfg.get("web_ui", {}).get("port", 5000)
        debug = cfg.get("web_ui", {}).get("debug", False)
    else:
        host, port, debug = "127.0.0.1", 5000, False

    print(f"  AI Journalist Studio running at http://{host}:{port}")
    app.run(host=host, port=port, debug=debug)
