from markdown import markdown
from docx import Document
from pathlib import Path
import datetime
import json
import re
import os


def _slugify(text):
    text = text.lower().strip()
    text = re.sub(r'[^\w\s-]', '', text)
    text = re.sub(r'[-\s]+', '-', text)
    return text


def publish_article(article_text: str, images: list, topic: str, article_type: str) -> dict:
    print("[5/5] Publishing article...")

    slug = _slugify(topic)
    today = datetime.date.today().isoformat()

    frontmatter = {
        "title": topic,
        "date": today,
        "type": article_type,
        "tags": [article_type, slug.replace("-", " ")]
    }
    yaml_lines = ["---"]
    for k, v in frontmatter.items():
        if isinstance(v, list):
            yaml_lines.append(f"{k}:")
            for item in v:
                yaml_lines.append(f"  - {item}")
        else:
            yaml_lines.append(f"{k}: {v}")
    yaml_lines.append("---")

    image_blocks = []
    for img in images:
        if isinstance(img, str):
            image_blocks.append(f"![Image]({img})\n")
        elif isinstance(img, dict):
            raw = img.get("filename") or img.get("path") or img.get("url", "")
            alt = img.get("caption") or img.get("alt", "") or "Image"
            project_root = str(Path(__file__).parent)
            rel = raw.replace(project_root + os.sep, "").replace(os.sep, "/")
            image_blocks.append(f"![{alt}](../{rel})\n")

    full_text = re.sub(r'!\[.*?\]\(\s*\)', '', article_text)
    if image_blocks:
        full_text = full_text + "\n\n" + "\n".join(image_blocks)

    md_content = "\n".join(yaml_lines) + "\n\n" + full_text

    articles_dir = Path("articles")
    exports_dir = Path("exports")
    articles_dir.mkdir(parents=True, exist_ok=True)
    exports_dir.mkdir(parents=True, exist_ok=True)

    md_path = articles_dir / f"{slug}.md"
    md_path.write_text(md_content, encoding="utf-8")
    print(f"  [SAVED] Markdown: {md_path}")

    html_path = exports_dir / f"{slug}.html"
    html_body = markdown(full_text, extensions=["extra"])
    html_doc = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{topic}</title>
<style>
body {{ font-family: 'Georgia', serif; max-width: 800px; margin: 2rem auto; padding: 0 1rem; line-height: 1.6; color: #222; }}
h1 {{ border-bottom: 2px solid #333; padding-bottom: 0.5rem; }}
img {{ max-width: 100%; height: auto; border-radius: 4px; }}
a {{ color: #1a73e8; }}
</style>
</head>
<body>
<h1>{topic}</h1>
<p><small>{today} &mdash; {article_type}</small></p>
{html_body}
</body>
</html>"""
    html_path.write_text(html_doc, encoding="utf-8")
    print(f"  [SAVED] HTML: {html_path}")

    docx_path = exports_dir / f"{slug}.docx"
    doc = Document()
    doc.add_heading(topic, level=1)
    doc.add_paragraph(f"{today} — {article_type}")
    for para in full_text.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        if para.startswith("#"):
            level = len(para.split()[0])
            text = para.lstrip("#").strip()
            doc.add_heading(text, level=min(level, 3))
        else:
            doc.add_paragraph(para)
    doc.save(str(docx_path))
    print(f"  [SAVED] DOCX: {docx_path}")

    pdf_path = None
    try:
        from weasyprint import HTML
        pdf_path = exports_dir / f"{slug}.pdf"
        HTML(string=html_doc).write_pdf(str(pdf_path))
        print(f"  [SAVED] PDF: {pdf_path}")
    except ImportError:
        print("  [SKIP] PDF — weasyprint not installed (pip install weasyprint)")
    except Exception as e:
        print(f"  [SKIP] PDF — conversion failed: {e}")

    mindmap_path = exports_dir / f"{slug}_mindmap.html"
    sections = [s.strip() for s in full_text.split("**") if s.strip() and len(s.strip()) > 10]
    nodes = [s for s in sections if not s.startswith("!") and not s.startswith("-")]
    mindmap_html = f"""<!DOCTYPE html>
<html lang="en">
<head><meta charset="utf-8"><title>{topic} - Mindmap</title>
<style>body{{margin:0;font-family:sans-serif;background:#1a1a2e;color:#e0e0e0;}}
#map{{width:100vw;height:100vh;}}
.node circle{{fill:#00d4aa;stroke:#00a080;stroke-width:2px;}}
.node text{{fill:#e0e0e0;font-size:12px;}}
.link{{stroke:#555;stroke-width:1.5px;}}</style>
</head>
<body><svg id="map"></svg>
<script src="https://d3js.org/d3.v7.min.js"></script>
<script>
var data={{name:"{topic[:50]}",children:[
{''.join(f'{{name:"{n[:60].replace(chr(34),"'")}"}},' for n in nodes[:15])}
]}};
var w=window.innerWidth,h=window.innerHeight;
var tree=d3.tree().size([h-100,w-200]);
var root=tree(d3.hierarchy(data));
var svg=d3.select("#map"),g=svg.append("g").attr("transform","translate(100,50)");
svg.attr("width",w).attr("height",h);
g.selectAll(".link").data(root.links()).enter().append("path").attr("class","link")
 .attr("d",d3.linkHorizontal().x(d=>d.y).y(d=>d.x));
var node=g.selectAll(".node").data(root.descendants()).enter().append("g")
 .attr("class","node").attr("transform",d=>`translate(${{d.y}},${{d.x}})`);
node.append("circle").attr("r",6);
node.append("text").attr("dy",-10).attr("x",d=>d.children?-10:10).attr("text-anchor",d=>d.children?"end":"start")
 .text(d=>d.data.name.slice(0,40));
</script></body></html>"""
    mindmap_path.write_text(mindmap_html, encoding="utf-8")
    print(f"  [SAVED] Mindmap: {mindmap_path}")

    result = {
        "markdown": str(md_path),
        "html": str(html_path),
        "docx": str(docx_path),
        "pdf": str(pdf_path) if pdf_path else None,
        "mindmap": str(mindmap_path),
    }

    print()
    print("  " + "-" * 40)
    print("  Publication complete")
    print("  " + "-" * 40)
    for fmt, path in result.items():
        if path:
            print(f"    {fmt.upper():>8}: {path}")
        else:
            print(f"    {fmt.upper():>8}: (skipped)")

    return result
